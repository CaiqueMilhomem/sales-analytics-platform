"""
Gera docs/estudo-completo.docx a partir do zero, etapa por etapa.

Por que regenerar do zero em vez de "abrir e acrescentar"? Editar um .docx
existente com python-docx tende a acumular inconsistência de estilo a cada
rodada. Mantendo o conteúdo de cada etapa como uma função aqui, o script é
a fonte da verdade: para registrar uma nova etapa, basta escrever uma nova
função `build_etapa_N` e chamá-la em main(). Rodar com: python gerar_estudo.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT_PATH = Path(__file__).parent / "estudo-completo.docx"

COR_TITULO = RGBColor(0x1F, 0x3A, 0x5F)
COR_CODIGO_FUNDO = RGBColor(0x2B, 0x2B, 0x2B)


def add_code_block(doc: Document, code: str) -> None:
    """Adiciona um bloco de código em fonte monoespaçada, parágrafo único."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_toc(doc: Document) -> None:
    """
    Insere um campo TOC (sumário) nativo do Word. python-docx não calcula
    números de página sozinho -- o Word preenche isso na primeira vez que o
    documento é aberto (ou com botão direito -> Atualizar campo).
    """
    doc.add_heading("Sumário", level=1)
    aviso = doc.add_paragraph()
    aviso.add_run(
        "(Clique com o botão direito sobre esta área e escolha "
        "\"Atualizar campo\" para gerar o sumário com números de página.)"
    ).italic = True

    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run()

    fld_inicio = OxmlElement("w:fldChar")
    fld_inicio.set(qn("w:fldCharType"), "begin")

    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_separador = OxmlElement("w:fldChar")
    fld_separador.set(qn("w:fldCharType"), "separate")

    fld_fim = OxmlElement("w:fldChar")
    fld_fim.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_inicio)
    r.append(instrucao)
    r.append(fld_separador)
    r.append(fld_fim)

    doc.add_page_break()


def add_titulo_capa(doc: Document) -> None:
    titulo = doc.add_heading("Estudo Completo — Sales Analytics Platform", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = COR_TITULO

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Documentação de estudo do fluxo SQL → automação em Python → "
        "endpoints FastAPI → integração com APIs externas"
    )
    run.italic = True
    doc.add_page_break()
    add_toc(doc)


def build_etapa_1(doc: Document) -> None:
    doc.add_heading("Etapa 1 — Fundação: Banco de Dados", level=1)

    # -------------------------------------------------------------------
    doc.add_heading("1. Cenário de negócio", level=2)
    doc.add_paragraph(
        "O projeto simula uma loja virtual (e-commerce) que precisa, "
        "rotineiramente, calcular métricas de vendas: faturamento por "
        "período, produtos mais vendidos, ticket médio por cliente e "
        "ranking de categorias. Esse cenário foi escolhido porque exige, "
        "de forma natural, as quatro tecnologias do fluxo de trabalho:"
    )
    itens = [
        ("SQL", "os dados vivem em um banco relacional e as regras de "
         "negócio (ex.: só contar vendas concluídas) devem ser calculadas "
         "uma única vez, no banco, via stored procedures — evitando que "
         "cada script ou relatório reimplemente a mesma lógica de forma "
         "divergente."),
        ("Automação em Python", "alguém precisa orquestrar a chamada "
         "dessas procedures periodicamente, tratar os resultados e "
         "preparar os dados para consumo — isso é rotina, não trabalho "
         "manual."),
        ("FastAPI", "o dashboard (ou qualquer outro consumidor) não deve "
         "acessar o banco diretamente; um endpoint HTTP expõe as métricas "
         "de forma controlada, documentada e testável."),
        ("Integração com API externa", "métricas internas ganham mais "
         "valor quando combinadas com contexto externo (ex.: cotação de "
         "moeda para vendas internacionais, dados de frete) — é o padrão "
         "mais comum de consolidação de dados para um dashboard real."),
    ]
    for nome, explicacao in itens:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{nome}: ")
        r.bold = True
        p.add_run(explicacao)

    # -------------------------------------------------------------------
    doc.add_heading("2. Estrutura de pastas", level=2)
    doc.add_paragraph(
        "O repositório separa claramente cada camada do fluxo, para que "
        "cada etapa do projeto tenha um lugar óbvio para crescer:"
    )
    add_code_block(
        doc,
        "sales-analytics-platform/\n"
        "├── sql/        # DDL, seed de dados e stored procedures\n"
        "├── scripts/    # Automações em Python (Etapa 2)\n"
        "├── api/        # Endpoints FastAPI (Etapa 3)\n"
        "├── data/       # Dados auxiliares/exportações\n"
        "├── docs/       # Este documento de estudo\n"
        "├── .env.example\n"
        "└── README.md",
    )

    # -------------------------------------------------------------------
    doc.add_heading("3. Modelagem do banco (sql/01_ddl.sql)", level=2)
    doc.add_paragraph(
        "Quatro tabelas modelam o domínio de vendas: clientes, produtos, "
        "vendas e itens_venda. O desenho segue uma decisão clássica de "
        "modelagem relacional: uma venda tem N itens (relação 1:N), e "
        "cada item referencia um produto — assim uma venda com 3 produtos "
        "diferentes vira 1 linha em vendas e 3 linhas em itens_venda, em "
        "vez de tentar espremer tudo em uma única tabela."
    )

    doc.add_heading("clientes e produtos", level=3)
    add_code_block(
        doc,
        "CREATE TABLE clientes (\n"
        "    id             SERIAL PRIMARY KEY,\n"
        "    nome           VARCHAR(150) NOT NULL,\n"
        "    email          VARCHAR(150) NOT NULL UNIQUE,\n"
        "    cidade         VARCHAR(100),\n"
        "    estado         CHAR(2),\n"
        "    data_cadastro  DATE NOT NULL DEFAULT CURRENT_DATE\n"
        ");\n\n"
        "CREATE TABLE produtos (\n"
        "    id              SERIAL PRIMARY KEY,\n"
        "    nome            VARCHAR(150) NOT NULL,\n"
        "    categoria       VARCHAR(80) NOT NULL,\n"
        "    preco_unitario  NUMERIC(10, 2) NOT NULL CHECK (preco_unitario > 0),\n"
        "    ativo           BOOLEAN NOT NULL DEFAULT TRUE\n"
        ");",
    )
    doc.add_paragraph(
        "SERIAL cria uma coluna inteira autoincrementável (o Postgres "
        "gerencia a sequência internamente) — é o jeito padrão de fazer "
        "chave primária numérica. O CHECK (preco_unitario > 0) impede, no "
        "próprio banco, que um preço inválido seja inserido — validar só "
        "na aplicação não é suficiente, porque qualquer acesso direto ao "
        "banco (um script, outra aplicação) poderia burlar essa regra."
    )

    doc.add_heading("vendas e itens_venda", level=3)
    add_code_block(
        doc,
        "CREATE TABLE vendas (\n"
        "    id          SERIAL PRIMARY KEY,\n"
        "    cliente_id  INTEGER NOT NULL REFERENCES clientes(id),\n"
        "    data_venda  TIMESTAMP NOT NULL DEFAULT NOW(),\n"
        "    status      VARCHAR(20) NOT NULL DEFAULT 'concluida'\n"
        "                CHECK (status IN ('concluida', 'pendente', 'cancelada'))\n"
        ");\n\n"
        "CREATE TABLE itens_venda (\n"
        "    id              SERIAL PRIMARY KEY,\n"
        "    venda_id        INTEGER NOT NULL REFERENCES vendas(id) ON DELETE CASCADE,\n"
        "    produto_id      INTEGER NOT NULL REFERENCES produtos(id),\n"
        "    quantidade      INTEGER NOT NULL CHECK (quantidade > 0),\n"
        "    preco_unitario  NUMERIC(10, 2) NOT NULL CHECK (preco_unitario > 0)\n"
        ");",
    )
    doc.add_paragraph(
        "REFERENCES clientes(id) é a chave estrangeira: garante que "
        "nenhuma venda aponte para um cliente que não existe. "
        "ON DELETE CASCADE em itens_venda.venda_id significa que, se uma "
        "venda for apagada, seus itens somem junto — faz sentido aqui "
        "porque um item de venda não tem existência própria fora da "
        "venda a que pertence."
    )
    doc.add_paragraph(
        "Ponto de atenção proposital: itens_venda guarda seu próprio "
        "preco_unitario, separado do preco_unitario em produtos. Isso é "
        "um \"snapshot\" do preço no momento da venda. Sem essa cópia, "
        "se o preço de um produto mudasse no catálogo, o faturamento "
        "histórico de vendas antigas mudaria retroativamente junto — o "
        "que não reflete a realidade (o cliente pagou o preço vigente na "
        "época da compra, não o preço de hoje)."
    )
    doc.add_paragraph(
        "Os índices em vendas.data_venda, vendas.cliente_id, "
        "itens_venda.venda_id e itens_venda.produto_id existem porque são "
        "exatamente as colunas usadas nos filtros e JOINs das stored "
        "procedures da seção 5 — sem índice, cada consulta faria uma "
        "varredura completa da tabela conforme o volume de dados cresce."
    )

    # -------------------------------------------------------------------
    doc.add_heading("4. Seed de dados (sql/02_dml_seed.sql)", level=2)
    doc.add_paragraph(
        "Popular 4000 vendas manualmente, linha por linha, não é "
        "realista. Em vez disso, o seed usa recursos do próprio SQL para "
        "gerar dados em volume:"
    )
    add_code_block(
        doc,
        "INSERT INTO vendas (cliente_id, data_venda, status)\n"
        "SELECT\n"
        "    1 + floor(random() * 300)::int,\n"
        "    NOW() - (random() * INTERVAL '730 days'),\n"
        "    CASE\n"
        "        WHEN random() < 0.92 THEN 'concluida'\n"
        "        WHEN random() < 0.97 THEN 'pendente'\n"
        "        ELSE 'cancelada'\n"
        "    END\n"
        "FROM generate_series(1, 4000);",
    )
    doc.add_paragraph(
        "generate_series(1, 4000) produz uma sequência de 4000 números "
        "que serve só para \"repetir a linha\" 4000 vezes — o valor em si "
        "não é usado. Para cada repetição, cliente_id é sorteado entre 1 "
        "e 300 (a faixa de IDs dos clientes recém-criados), a data cai em "
        "algum ponto dos últimos 730 dias, e o status é sorteado com "
        "pesos (92% concluída, 5% pendente, 3% cancelada) para simular "
        "uma operação real, onde a maioria das vendas fecha."
    )
    add_code_block(
        doc,
        "INSERT INTO itens_venda (venda_id, produto_id, quantidade, preco_unitario)\n"
        "SELECT v.id, item.produto_id, item.quantidade, prod.preco_unitario\n"
        "FROM vendas v\n"
        "CROSS JOIN LATERAL (\n"
        "    SELECT\n"
        "        1 + floor(random() * 40)::int AS produto_id,\n"
        "        1 + floor(random() * 5)::int  AS quantidade\n"
        "    FROM generate_series(1, 1 + floor(random() * 4)::int)\n"
        ") item\n"
        "JOIN produtos prod ON prod.id = item.produto_id;",
    )
    doc.add_paragraph(
        "Este é o bloco mais denso do seed. CROSS JOIN LATERAL permite "
        "que a subconsulta à direita seja reavaliada para cada linha de "
        "vendas — sem LATERAL, a subconsulta seria calculada uma única "
        "vez e repetida para todas as vendas, o que não serve aqui, "
        "porque cada venda precisa de uma quantidade de itens diferente "
        "e sorteada de novo. O generate_series(1, 1 + floor(random()*4)) "
        "interno decide, para aquela venda específica, quantos itens ela "
        "vai ter (entre 1 e 4). O resultado final: ~4000 vendas geram "
        "aproximadamente 10 mil linhas em itens_venda."
    )

    # -------------------------------------------------------------------
    doc.add_heading("5. Stored procedures (sql/03_stored_procedures.sql)", level=2)
    doc.add_paragraph(
        "Nota de vocabulário importante: no PostgreSQL, uma rotina que "
        "PRECISA devolver um conjunto de linhas (uma tabela de "
        "resultado) é criada como FUNCTION, não PROCEDURE. PROCEDURE "
        "existe desde o Postgres 11 e é chamada com CALL, mas serve "
        "principalmente para rotinas que executam ações (INSERT/UPDATE/"
        "DELETE, controle de transação) e não têm um jeito direto de "
        "devolver várias linhas como resultado de consulta. Por isso as "
        "4 rotinas abaixo são FUNCTIONs com RETURNS TABLE + RETURN "
        "QUERY, e são testadas com SELECT * FROM nome(...), não CALL. "
        "Em bancos como SQL Server ou Oracle, o termo \"stored "
        "procedure\" costuma ser usado de forma genérica para esse tipo "
        "de rotina — vale saber reconhecer a diferença de nomenclatura "
        "entre bancos."
    )
    doc.add_paragraph(
        "Regra de negócio comum às quatro: somente vendas com "
        "status = 'concluida' entram no cálculo de faturamento. Vendas "
        "pendentes ou canceladas não representam receita real."
    )

    doc.add_heading("fn_resumo_vendas_periodo", level=3)
    add_code_block(
        doc,
        "CREATE OR REPLACE FUNCTION fn_resumo_vendas_periodo(\n"
        "    p_data_inicio DATE, p_data_fim DATE\n"
        ")\n"
        "RETURNS TABLE (total_vendas BIGINT, faturamento_total NUMERIC,\n"
        "               ticket_medio NUMERIC, total_itens BIGINT)\n"
        "LANGUAGE plpgsql AS $$\n"
        "BEGIN\n"
        "    IF p_data_inicio > p_data_fim THEN\n"
        "        RAISE EXCEPTION 'data_inicio (%) nao pode ser maior que data_fim (%)',\n"
        "            p_data_inicio, p_data_fim;\n"
        "    END IF;\n\n"
        "    RETURN QUERY\n"
        "    SELECT\n"
        "        COUNT(DISTINCT v.id)::BIGINT,\n"
        "        COALESCE(SUM(iv.quantidade * iv.preco_unitario), 0)::NUMERIC,\n"
        "        COALESCE(SUM(iv.quantidade * iv.preco_unitario)\n"
        "                 / NULLIF(COUNT(DISTINCT v.id), 0), 0)::NUMERIC,\n"
        "        COALESCE(SUM(iv.quantidade), 0)::BIGINT\n"
        "    FROM vendas v\n"
        "    JOIN itens_venda iv ON iv.venda_id = v.id\n"
        "    WHERE v.status = 'concluida'\n"
        "      AND v.data_venda::date BETWEEN p_data_inicio AND p_data_fim;\n"
        "END;\n$$;",
    )
    doc.add_paragraph(
        "A validação IF p_data_inicio > p_data_fim ... RAISE EXCEPTION é "
        "o tratamento de erro: em vez de devolver um resultado sem "
        "sentido (ou vazio, silenciosamente), a função interrompe a "
        "execução com uma mensagem clara. NULLIF(COUNT(...), 0) evita "
        "divisão por zero no cálculo do ticket médio quando não há "
        "nenhuma venda no período — sem isso, o Postgres lançaria um "
        "erro de \"division by zero\" em vez de simplesmente devolver 0."
    )

    doc.add_heading("fn_top_produtos_faturamento, fn_faturamento_ticket_medio_cliente, fn_vendas_categoria_ranking", level=3)
    doc.add_paragraph(
        "As outras três seguem a mesma estrutura (validação de "
        "parâmetros + RETURN QUERY), variando o agrupamento:"
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("fn_top_produtos_faturamento: ")
    r.bold = True
    p.add_run(
        "agrupa por produto, soma quantidade e faturamento, ordena "
        "decrescente e usa LIMIT p_limite para trazer só o top N. "
        "Também valida que p_limite seja maior que zero."
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("fn_faturamento_ticket_medio_cliente: ")
    r.bold = True
    p.add_run(
        "agrupa por cliente. Tem um parâmetro opcional p_cliente_id "
        "(DEFAULT NULL) — quando informado, filtra um único cliente; "
        "quando omitido, devolve o ranking de todos. A condição "
        "(p_cliente_id IS NULL OR c.id = p_cliente_id) é o padrão usual "
        "para \"filtro opcional\" em SQL."
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("fn_vendas_categoria_ranking: ")
    r.bold = True
    p.add_run(
        "usa uma window function, RANK() OVER (ORDER BY faturamento_total "
        "DESC), para numerar as categorias por faturamento sem precisar "
        "de uma segunda consulta. Diferente de GROUP BY sozinho, uma "
        "window function calcula o agregado (faturamento por categoria, "
        "na subconsulta) e ainda permite atribuir um rank a cada linha "
        "do resultado, mantendo uma linha por categoria."
    )

    # -------------------------------------------------------------------
    doc.add_heading("6. Executando no Neon/Supabase", level=2)
    doc.add_paragraph(
        "Passo a passo resumido (detalhes completos no README.md do "
        "projeto):"
    )
    passos = [
        "Criar conta e projeto gratuito no Neon (neon.tech) ou Supabase "
        "(supabase.com).",
        "Abrir o SQL Editor do navegador do serviço escolhido.",
        "Colar e executar sql/01_ddl.sql (cria as 4 tabelas).",
        "Colar e executar sql/02_dml_seed.sql (popula os dados de exemplo).",
        "Colar e executar sql/03_stored_procedures.sql (cria as 4 funções).",
        "Copiar a connection string do painel do serviço e colar em "
        "DATABASE_URL no arquivo .env (copiado a partir de .env.example).",
        "Testar cada função com SELECT * FROM nome_da_funcao(...) "
        "diretamente no SQL Editor.",
    ]
    for passo in passos:
        doc.add_paragraph(passo, style="List Number")

    # -------------------------------------------------------------------
    doc.add_heading("7. Conceitos-chave desta etapa", level=2)
    conceitos = [
        ("SERIAL", "tipo que cria automaticamente uma sequência e usa "
         "seus valores como padrão da coluna — a forma padrão de "
         "chave primária autoincrementável no Postgres."),
        ("FOREIGN KEY / REFERENCES", "garante integridade referencial: "
         "impede que uma linha aponte para um registro que não existe "
         "em outra tabela."),
        ("CHECK constraint", "regra de validação garantida pelo próprio "
         "banco, independente de qual aplicação está inserindo dados."),
        ("Snapshot de preço", "copiar um valor (preço) no momento do "
         "evento (venda) em vez de sempre referenciar o valor atual, "
         "para preservar o histórico correto."),
        ("generate_series", "função que gera uma sequência de números "
         "(ou datas), útil para criar N linhas repetidas em massa."),
        ("LATERAL join", "permite que uma subconsulta à direita do JOIN "
         "referencie colunas da linha à esquerda e seja reavaliada por "
         "linha — necessário quando cada linha externa precisa de um "
         "resultado calculado de forma independente."),
        ("PL/pgSQL", "linguagem procedural do Postgres usada para "
         "escrever functions/procedures com lógica (IF, RAISE "
         "EXCEPTION, variáveis), além de SQL puro."),
        ("RETURN QUERY", "dentro de uma function PL/pgSQL, executa uma "
         "consulta e devolve seu resultado como o result set da "
         "função."),
        ("FUNCTION vs PROCEDURE", "no Postgres, FUNCTION devolve valor/"
         "tabela e é chamada com SELECT; PROCEDURE é chamada com CALL e "
         "não devolve result set da mesma forma — a nomenclatura "
         "genérica \"stored procedure\" usada em outros bancos não bate "
         "1:1 com a sintaxe do Postgres."),
        ("Window function (RANK)", "calcula um valor por linha (como "
         "uma posição em um ranking) considerando um conjunto de linhas "
         "relacionado, sem colapsar o resultado em uma única linha como "
         "GROUP BY faria."),
    ]
    for termo, explicacao in conceitos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_page_break()


def build_etapa_2(doc: Document) -> None:
    doc.add_heading("Etapa 2 — Automação em Python", level=1)

    # -------------------------------------------------------------------
    doc.add_heading("1. Visão geral", level=2)
    doc.add_paragraph(
        "A pasta scripts/ é a camada que transforma as stored functions da "
        "Etapa 1 em relatórios prontos para consumo: conecta no Postgres, "
        "chama cada função SQL, converte o resultado em DataFrame do "
        "pandas, salva em CSV/JSON e monta um resumo executivo. O código "
        "é dividido em três arquivos, cada um com uma responsabilidade "
        "clara:"
    )
    responsabilidades = [
        ("database.py", "conexão com o banco (context manager) e execução "
         "genérica de SELECT * FROM fn_...(...)."),
        ("metricas.py", "uma função por stored function, cada uma já "
         "devolvendo um DataFrame."),
        ("gerar_relatorios.py", "o CLI: resolve o período pedido, chama as "
         "métricas, salva os arquivos e monta o resumo executivo."),
    ]
    for nome, explicacao in responsabilidades:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{nome}: ")
        r.bold = True
        p.add_run(explicacao)
    doc.add_paragraph(
        "Essa separação existe para que cada camada possa ser testada e "
        "entendida isoladamente: quem quer saber como a conexão é aberta "
        "olha só database.py; quem quer saber que métricas existem olha "
        "metricas.py; a orquestração (CLI, período, arquivos) fica só em "
        "gerar_relatorios.py."
    )

    # -------------------------------------------------------------------
    doc.add_heading("2. Conexão com o banco (scripts/database.py)", level=2)
    add_code_block(
        doc,
        "@contextmanager\n"
        "def conexao():\n"
        "    database_url = os.getenv('DATABASE_URL')\n"
        "    if not database_url:\n"
        "        raise ConfiguracaoAusente(...)\n\n"
        "    conn = psycopg2.connect(database_url)\n"
        "    try:\n"
        "        yield conn\n"
        "    except Exception:\n"
        "        conn.rollback()\n"
        "        raise\n"
        "    finally:\n"
        "        conn.close()",
    )
    doc.add_paragraph(
        "@contextmanager (do módulo contextlib) transforma essa função "
        "geradora em algo usável com `with conexao() as conn:`. O ponto "
        "central: o bloco finally garante que conn.close() roda sempre — "
        "sucesso, erro de SQL, ou até um Ctrl+C no meio do caminho — "
        "então nenhuma chamada às métricas corre risco de vazar conexão "
        "aberta. O except antes do finally faz rollback caso alguma "
        "operação tenha ficado pendente numa transação, antes de fechar."
    )
    doc.add_paragraph(
        "Duas decisões técnicas não óbvias, registradas como comentário "
        "no código e explicadas aqui com mais espaço:"
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Conversão NUMERIC → float: ")
    r.bold = True
    p.add_run(
        "por padrão, o psycopg2 devolve colunas NUMERIC do Postgres como "
        "decimal.Decimal — tipo que o pandas consegue guardar, mas que o "
        "método to_json() não sabe serializar (gera TypeError). A solução "
        "usada foi registrar um conversor global "
        "(psycopg2.extensions.register_type) que troca NUMERIC por float "
        "assim que os dados chegam do banco. É um trade-off consciente: "
        "float perde um pouco de precisão decimal em casos extremos, "
        "aceitável para relatório, mas não deveria ser usado em um "
        "sistema que faça, por exemplo, fechamento contábil."
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("SQL composto com psycopg2.sql: ")
    r.bold = True
    p.add_run(
        "executar_funcao() monta o nome da função dentro da query usando "
        "sql.Identifier(nome_funcao) em vez de simplesmente colar a "
        "string com f-string. Hoje esse nome sempre vem de um literal "
        "fixo dentro de metricas.py, então o risco de SQL injection é "
        "teórico — mas é o tipo de hábito que evita que o problema exista "
        "no dia em que alguém decidir aceitar esse nome vindo de um "
        "parâmetro externo (ex.: um endpoint da API na Etapa 3)."
    )

    # -------------------------------------------------------------------
    doc.add_heading("3. Camada de métricas (scripts/metricas.py)", level=2)
    add_code_block(
        doc,
        "def top_produtos_faturamento(data_inicio, data_fim, limite=10):\n"
        "    colunas, linhas = executar_funcao(\n"
        "        'fn_top_produtos_faturamento', (data_inicio, data_fim, limite)\n"
        "    )\n"
        "    return pd.DataFrame(linhas, columns=colunas)",
    )
    doc.add_paragraph(
        "Cada uma das 4 funções segue exatamente este padrão: chama "
        "executar_funcao() com o nome da stored function e os parâmetros "
        "na mesma ordem definida no SQL, e devolve um pd.DataFrame usando "
        "as colunas retornadas pelo cursor (cur.description) — assim, se "
        "uma coluna for renomeada na function SQL, o DataFrame acompanha "
        "automaticamente, sem precisar editar essa função Python."
    )

    # -------------------------------------------------------------------
    doc.add_heading("4. CLI e orquestração (scripts/gerar_relatorios.py)", level=2)
    doc.add_heading("Resolução do período (--periodo)", level=3)
    add_code_block(
        doc,
        "def resolver_periodo(valor: str) -> Periodo:\n"
        "    hoje = date.today()\n"
        "    if ':' in valor:\n"
        "        inicio_str, fim_str = valor.split(':', 1)\n"
        "        ...\n"
        "    if valor == 'mes-atual':\n"
        "        return Periodo(hoje.replace(day=1), hoje)\n"
        "    if valor.startswith('ultimos-') and valor.endswith('d'):\n"
        "        dias = int(valor[len('ultimos-'):-1])\n"
        "        return Periodo(hoje - timedelta(days=dias), hoje)\n"
        "    raise ValueError(...)",
    )
    doc.add_paragraph(
        "Periodo é um @dataclass (só dois campos, inicio e fim) — usado "
        "aqui em vez de uma tupla solta porque `periodo.inicio` é mais "
        "legível que `periodo[0]` no resto do código. resolver_periodo "
        "aceita tanto atalhos pensados para agendamento automático "
        "(ultimos-30d, mes-atual, ano-atual) quanto um intervalo explícito "
        "(AAAA-MM-DD:AAAA-MM-DD) para reprocessar um período específico "
        "manualmente. Uma entrada que não bate com nenhum padrão levanta "
        "ValueError com uma mensagem que já explica o formato esperado."
    )

    doc.add_heading("Resumo executivo e exportação", level=3)
    add_code_block(
        doc,
        "def salvar(df, caminho_base, formatos):\n"
        "    if 'csv' in formatos:\n"
        "        df.to_csv(caminho_base.with_suffix('.csv'), index=False)\n"
        "    if 'json' in formatos:\n"
        "        df.to_json(caminho_base.with_suffix('.json'),\n"
        "                    orient='records', indent=2, force_ascii=False)",
    )
    doc.add_paragraph(
        "orient='records' grava o JSON como uma lista de objetos "
        "(um por linha do DataFrame), o formato mais fácil de consumir "
        "por outro programa. force_ascii=False mantém acentos legíveis "
        "no arquivo em vez de escapá-los em \\uXXXX."
    )
    doc.add_paragraph(
        "montar_resumo_executivo() pega a primeira (e única) linha de "
        "fn_resumo_vendas_periodo, a primeira linha de "
        "fn_top_produtos_faturamento (produto de maior faturamento) e a "
        "primeira linha de fn_vendas_categoria_ranking (categoria "
        "líder, já que a função SQL devolve ordenado por ranking), e "
        "consolida tudo em um único dicionário — esse é o "
        "resumo_executivo_*.json, pensado para ser o arquivo que alguém "
        "abre primeiro para ter uma visão geral, sem precisar olhar os "
        "quatro relatórios detalhados."
    )

    doc.add_heading("Tratamento de erro e códigos de saída", level=3)
    doc.add_paragraph(
        "executar() nunca deixa uma exceção subir crua até o topo do "
        "programa: ValueError de período inválido vira código de saída "
        "2, ConfiguracaoAusente (variável de ambiente faltando) e "
        "qualquer outro erro de banco viram código 1, e só quando tudo "
        "funciona o programa devolve 0. Essa distinção importa porque um "
        "agendador (cron, Task Scheduler) decide se deve alertar alguém "
        "só olhando o código de saída do processo — sys.exit(codigo) é a "
        "forma padrão de comunicar isso para quem chamou o script."
    )

    doc.add_heading("Argumentos de linha de comando (argparse)", level=3)
    doc.add_paragraph(
        "argparse.ArgumentParser define cada flag com tipo, valor padrão "
        "e texto de ajuda (visível com --help). --formato usa "
        "choices=['csv', 'json', 'ambos'] para que um valor fora dessa "
        "lista já seja rejeitado pelo próprio argparse, com mensagem de "
        "erro padronizada, antes mesmo do código do projeto rodar."
    )

    # -------------------------------------------------------------------
    doc.add_heading("5. Instalação sem privilégio de administrador", level=2)
    doc.add_paragraph(
        "Como não é possível instalar nada globalmente na máquina da "
        "empresa, o requirements.txt é instalado dentro de um ambiente "
        "virtual (venv) — uma cópia isolada do Python que vive dentro da "
        "própria pasta do projeto, sem exigir permissão elevada:"
    )
    add_code_block(
        doc,
        "python -m venv .venv\n"
        ".venv\\Scripts\\Activate.ps1   # PowerShell\n"
        "pip install -r requirements.txt",
    )
    doc.add_paragraph(
        "Se até a criação do venv for bloqueada pela política da "
        "máquina, a alternativa é pip install --user -r requirements.txt, "
        "que instala os pacotes só na área do usuário atual, sem tocar "
        "na instalação Python compartilhada do sistema."
    )
    doc.add_paragraph(
        "Dependências (requirements.txt): psycopg2-binary (driver "
        "Postgres — a variante \"binary\" já vem compilada, evitando a "
        "necessidade de um compilador C instalado, o que é importante "
        "numa máquina sem ferramentas de build), pandas (transformação "
        "dos dados) e python-dotenv (carrega o .env para variáveis de "
        "ambiente do processo)."
    )

    # -------------------------------------------------------------------
    doc.add_heading("6. Conceitos-chave desta etapa", level=2)
    conceitos = [
        ("Context manager (@contextmanager)", "padrão que garante que um "
         "recurso (aqui, a conexão com o banco) seja liberado mesmo se o "
         "código dentro do `with` lançar uma exceção."),
        ("psycopg2.sql (composição segura de SQL)", "monta identificadores "
         "e placeholders de forma que o driver escapa/valida cada parte, "
         "em vez de concatenar strings manualmente — evita SQL injection "
         "mesmo quando o valor de entrada é, hoje, controlado."),
        ("Trade-off Decimal vs float", "decisão consciente de perder um "
         "pouco de precisão decimal em troca de compatibilidade direta "
         "com pandas/JSON — válida para relatório, não para cálculo "
         "financeiro crítico."),
        ("dataclass", "forma enxuta de criar uma classe que só guarda "
         "dados (aqui, Periodo com inicio/fim), sem precisar escrever "
         "__init__ manualmente."),
        ("argparse", "módulo padrão do Python para ler argumentos de "
         "linha de comando, gerar --help automaticamente e validar tipos/"
         "opções antes do código da aplicação rodar."),
        ("logging vs print", "logging.info/error grava timestamp, nível "
         "de severidade e permite ligar/desligar detalhamento "
         "(--verbose) sem mudar o código — diferente de print, que não "
         "distingue nível de log nem é fácil de redirecionar em produção."),
        ("Código de saída (exit code)", "convenção usada por processos de "
         "linha de comando para informar sucesso (0) ou falha (diferente "
         "de 0) a quem os chamou — é como um cron/agendador decide "
         "disparar um alerta sem precisar interpretar o texto do log."),
        ("DataFrame.to_csv / to_json", "métodos do pandas que exportam um "
         "DataFrame diretamente para arquivo, sem precisar escrever o "
         "parsing/formatação manualmente."),
    ]
    for termo, explicacao in conceitos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_page_break()


def build_etapa_3(doc: Document) -> None:
    doc.add_heading("Etapa 3 — API com FastAPI", level=1)

    # -------------------------------------------------------------------
    doc.add_heading("1. Visão geral", level=2)
    doc.add_paragraph(
        "A pasta api/ expõe as mesmas métricas da Etapa 2 como endpoints "
        "HTTP, para que um dashboard (ou qualquer outro consumidor) não "
        "precise ter acesso direto ao Postgres — ele conversa só com a "
        "API, que decide o que expor, valida entrada e formata a saída."
    )
    tabela = [
        ("main.py", "instância da app, CORS, handlers de erro e o health check."),
        ("routers/vendas.py", "/vendas/resumo, /vendas/top-produtos, /vendas/por-categoria."),
        ("routers/clientes.py", "/clientes/top."),
        ("schemas.py", "modelos Pydantic das respostas (o \"contrato\" de cada rota)."),
        ("dependencies.py", "validação do parâmetro periodo, compartilhada por todas as rotas."),
        ("data_access.py", "ponte para scripts/database.py e scripts/metricas.py."),
    ]
    for nome, explicacao in tabela:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{nome}: ")
        r.bold = True
        p.add_run(explicacao)

    # -------------------------------------------------------------------
    doc.add_heading("2. Reaproveitando a Etapa 2 (api/data_access.py)", level=2)
    doc.add_paragraph(
        "scripts/ foi escrito na Etapa 2 para rodar como scripts soltos "
        "(database.py e metricas.py se importam um ao outro com `import "
        "database`, sem formar um pacote Python de verdade). Para a API "
        "reaproveitar esse código sem copiá-lo, data_access.py insere a "
        "pasta scripts/ no sys.path uma única vez, na inicialização, e "
        "reexporta o que as rotas precisam:"
    )
    add_code_block(
        doc,
        "_SCRIPTS_DIR = Path(__file__).resolve().parent.parent / 'scripts'\n"
        "if str(_SCRIPTS_DIR) not in sys.path:\n"
        "    sys.path.insert(0, str(_SCRIPTS_DIR))\n\n"
        "import metricas\n"
        "from database import ConfiguracaoAusente\n"
        "from periodo import Periodo, resolver_periodo",
    )
    doc.add_paragraph(
        "É um atalho pragmático, não a única forma de resolver isso — em "
        "um projeto maior, o caminho mais \"correto\" seria transformar "
        "scripts/ em um pacote instalável (pyproject.toml + pip install "
        "-e .) e importar normalmente. Para o tamanho atual do projeto, "
        "manipular o sys.path em um único ponto de entrada é mais simples "
        "e igualmente explícito — o comentário no código já registra essa "
        "escolha."
    )
    doc.add_paragraph(
        "Nenhuma rota importa database.py ou metricas.py diretamente: "
        "todas passam por api/data_access.py. Isso significa que, se um "
        "dia a forma de acessar o banco mudar, só esse arquivo precisa "
        "mudar."
    )

    # -------------------------------------------------------------------
    doc.add_heading("3. Refatoração: scripts/periodo.py", level=2)
    doc.add_paragraph(
        "A validação do parâmetro periodo (ultimos-Nd, mes-atual, "
        "ano-atual, AAAA-MM-DD:AAAA-MM-DD) já existia no CLI da Etapa 2, "
        "dentro de gerar_relatorios.py. Em vez de colar essa mesma lógica "
        "dentro da API, ela foi extraída para scripts/periodo.py — agora "
        "tanto o CLI quanto a API importam a mesma função "
        "resolver_periodo(). Essa é a razão prática por trás da regra "
        "\"regra de negócio em um só lugar\", aplicada de novo aqui: sem "
        "essa extração, um ajuste futuro no formato aceito por --periodo "
        "exigiria lembrar de mudar dois lugares — e mais cedo ou mais "
        "tarde alguém esqueceria um deles."
    )

    # -------------------------------------------------------------------
    doc.add_heading("4. Schemas Pydantic e response_model", level=2)
    add_code_block(
        doc,
        "class ProdutoTop(BaseModel):\n"
        "    produto_id: int\n"
        "    produto_nome: str\n"
        "    categoria: str\n"
        "    quantidade_vendida: int\n"
        "    faturamento: float",
    )
    doc.add_paragraph(
        "Cada schema em api/schemas.py descreve exatamente os campos que "
        "uma rota promete devolver. Ao declarar response_model=ProdutoTop "
        "(ou list[ProdutoTop]) na rota, o FastAPI valida o dicionário "
        "devolvido pela função contra esse modelo antes de mandar a "
        "resposta — se um campo estiver faltando ou com tipo errado, o "
        "erro aparece imediatamente (500, com o problema registrado no "
        "log) em vez de um JSON malformado silenciosamente chegando ao "
        "cliente. É também esse schema que gera a documentação automática "
        "em /docs, com o formato de cada resposta já preenchido."
    )

    # -------------------------------------------------------------------
    doc.add_heading("5. Dependency injection (api/dependencies.py)", level=2)
    add_code_block(
        doc,
        "def obter_periodo(periodo: str = 'ultimos-30d') -> Periodo:\n"
        "    try:\n"
        "        return resolver_periodo(periodo)\n"
        "    except ValueError as exc:\n"
        "        raise HTTPException(status_code=422, detail=str(exc)) from exc\n\n"
        "PeriodoDep = Annotated[Periodo, Depends(obter_periodo)]",
    )
    doc.add_paragraph(
        "Depends() é o mecanismo de injeção de dependência do FastAPI: "
        "declarar `periodo: PeriodoDep` como parâmetro de uma rota faz o "
        "framework chamar obter_periodo() automaticamente antes da rota "
        "rodar, usando o valor de ?periodo=... da URL. PeriodoDep é só um "
        "apelido (Annotated[Periodo, Depends(obter_periodo)]) para não "
        "repetir essa anotação em cada uma das 4 rotas. Aqui está o único "
        "lugar da API que converte um erro de validação (ValueError) em "
        "HTTPException — as rotas em si não precisam se preocupar com "
        "isso, porque o parâmetro já chega validado."
    )

    # -------------------------------------------------------------------
    doc.add_heading("6. As rotas (api/routers/)", level=2)
    add_code_block(
        doc,
        "@router.get('/resumo', response_model=ResumoVendas)\n"
        "def resumo_vendas(periodo: PeriodoDep):\n"
        "    df = metricas.resumo_vendas_periodo(periodo.inicio, periodo.fim)\n"
        "    linha = df.iloc[0].to_dict()\n"
        "    return {'periodo': {'inicio': periodo.inicio, 'fim': periodo.fim}, **linha}\n\n"
        "@router.get('/top-produtos', response_model=list[ProdutoTop])\n"
        "def top_produtos(\n"
        "    periodo: PeriodoDep,\n"
        "    limite: int = Query(10, ge=1, le=100),\n"
        "):\n"
        "    df = metricas.top_produtos_faturamento(periodo.inicio, periodo.fim, limite)\n"
        "    return df.to_dict(orient='records')",
    )
    doc.add_paragraph(
        "Query(10, ge=1, le=100) faz o FastAPI validar limite antes mesmo "
        "da rota executar: um valor fora de 1-100 já devolve 422 "
        "automaticamente, sem precisar de um if manual dentro da função. "
        "df.to_dict(orient='records') transforma o DataFrame em uma lista "
        "de dicionários — formato que response_model=list[ProdutoTop] "
        "sabe validar e serializar."
    )
    doc.add_paragraph(
        "Ponto de atenção deliberado: todas as rotas usam def comum, não "
        "async def. metricas.* chama psycopg2, que é uma biblioteca "
        "síncrona (bloqueante) — se uma rota async def fizesse essa "
        "chamada diretamente, ela travaria o event loop inteiro do "
        "servidor até a consulta terminar, inclusive para outras "
        "requisições simultâneas. O FastAPI já executa rotas síncronas "
        "automaticamente numa threadpool, então usar def aqui é a forma "
        "correta de lidar com uma biblioteca de banco sem suporte a "
        "async."
    )
    doc.add_paragraph(
        "/clientes/top merece uma nota: fn_faturamento_ticket_medio_"
        "cliente (Etapa 1) não tem parâmetro de limite — ela sempre "
        "devolve todos os clientes ordenados por faturamento decrescente. "
        "O corte de \"top N\" é feito depois, em Python, com "
        "df.head(limite). Dado o volume do projeto (no máximo ~300 "
        "clientes), buscar tudo e cortar depois é mais simples do que "
        "criar uma variação da função SQL só para isso."
    )

    # -------------------------------------------------------------------
    doc.add_heading("7. Tratamento de erros centralizado (api/main.py)", level=2)
    add_code_block(
        doc,
        "@app.exception_handler(ConfiguracaoAusente)\n"
        "def _configuracao_ausente_handler(request, exc):\n"
        "    return JSONResponse(status_code=500, content={'detail': str(exc)})\n\n"
        "@app.exception_handler(Exception)\n"
        "def _erro_inesperado_handler(request, exc):\n"
        "    logger.exception('Erro nao tratado ao processar %s', request.url.path)\n"
        "    return JSONResponse(status_code=500, content={'detail': 'Erro interno ao consultar os dados.'})",
    )
    doc.add_paragraph(
        "Em vez de repetir um try/except em cada uma das 4 rotas, dois "
        "exception_handler cobrem toda a aplicação: um específico para "
        "ConfiguracaoAusente (DATABASE_URL faltando, devolve a mensagem "
        "exata do erro — é seguro expor, não vaza dado sensível) e um "
        "genérico para qualquer outra exceção (erro de SQL, timeout de "
        "conexão etc.), que loga o traceback completo no servidor mas "
        "devolve uma mensagem curta e genérica ao cliente, sem detalhes "
        "internos. O handler genérico não interfere no HTTPException "
        "lançado por obter_periodo(): o FastAPI já registra um handler "
        "próprio para HTTPException e ele tem prioridade, porque é mais "
        "específico na hierarquia de classes do que Exception."
    )
    doc.add_paragraph(
        "Resumindo a distinção: HTTPException é usado onde o próprio "
        "código da rota/dependência sabe exatamente o que deu errado e "
        "que status devolver (422 para parâmetro inválido); os "
        "exception_handler cobrem o que pode falhar em qualquer rota sem "
        "aviso prévio (banco fora do ar, configuração ausente)."
    )

    # -------------------------------------------------------------------
    doc.add_heading("8. CORS", level=2)
    add_code_block(
        doc,
        "app.add_middleware(\n"
        "    CORSMiddleware,\n"
        "    allow_origins=['*'] if origens == '*' else [o.strip() for o in origens.split(',')],\n"
        "    allow_methods=['GET'],\n"
        "    allow_headers=['*'],\n"
        ")",
    )
    doc.add_paragraph(
        "CORS (Cross-Origin Resource Sharing) é a política do navegador "
        "que decide se um site em um domínio pode chamar uma API em "
        "outro. Sem esse middleware, um dashboard servido em outro "
        "endereço seria bloqueado pelo próprio navegador ao tentar "
        "consumir a API. CORS_ORIGINS é lido do .env: \"*\" (padrão) "
        "libera qualquer origem, prático para desenvolvimento local; em "
        "produção, o certo é restringir à lista real de domínios que "
        "consomem a API. allow_methods fica limitado a GET porque é tudo "
        "que a API expõe hoje — menos superfície liberada por padrão."
    )

    # -------------------------------------------------------------------
    doc.add_heading("9. Rodando e testando", level=2)
    doc.add_paragraph(
        "A partir da raiz do projeto (não de dentro de api/), com o "
        "ambiente virtual da Etapa 2 já ativado (requirements.txt agora "
        "inclui fastapi e uvicorn[standard]):"
    )
    add_code_block(doc, "uvicorn api.main:app --reload")
    doc.add_paragraph(
        "Rodar a partir da raiz importa: é isso que faz o Python "
        "enxergar api/ como um pacote (api.main:app) e resolver os "
        "imports absolutos usados nos routers (from api.data_access "
        "import metricas). A documentação interativa fica disponível em "
        "http://127.0.0.1:8000/docs (Swagger UI, gerado automaticamente "
        "a partir dos schemas e das rotas) — cada endpoint pode ser "
        "testado diretamente por ali, com formulário de parâmetros."
    )
    doc.add_paragraph(
        "Validação feita nesta etapa (sem um Postgres real disponível): "
        "subi a aplicação de verdade com uvicorn e também via "
        "fastapi.testclient.TestClient, e confirmei que o health check "
        "responde sem tocar no banco, que periodo/limite inválidos "
        "devolvem 422 antes de qualquer consulta, e que, sem "
        "DATABASE_URL configurada, as quatro rotas de métricas devolvem "
        "500 com mensagem clara em vez de travar o processo — exatamente "
        "o comportamento que os exception_handler foram desenhados para "
        "garantir."
    )

    # -------------------------------------------------------------------
    doc.add_heading("10. Conceitos-chave desta etapa", level=2)
    conceitos = [
        ("FastAPI + Pydantic", "framework web que usa type hints Python "
         "para validar entrada, serializar saída e gerar documentação "
         "automática (Swagger) a partir do mesmo código, sem escrever a "
         "documentação à parte."),
        ("APIRouter", "agrupa rotas relacionadas (vendas, clientes) em "
         "módulos separados, incluídos na app principal com "
         "app.include_router() — evita um único main.py gigante."),
        ("response_model", "contrato de saída de uma rota: valida e "
         "filtra o retorno contra um schema Pydantic antes de enviar a "
         "resposta."),
        ("Dependency injection (Depends)", "mecanismo do FastAPI para "
         "reaproveitar lógica (aqui, validação de periodo) entre "
         "múltiplas rotas, declarada uma vez e injetada automaticamente "
         "como parâmetro."),
        ("Annotated", "forma moderna de combinar o tipo de um parâmetro "
         "(Periodo) com metadados extras (Depends(obter_periodo)) numa "
         "única anotação, usada tanto em Query quanto em Depends."),
        ("exception_handler", "captura centralizada de um tipo de "
         "exceção para toda a aplicação, em vez de try/except repetido "
         "em cada rota."),
        ("Rotas síncronas em threadpool", "o FastAPI roda automaticamente "
         "toda rota declarada com def (não async def) numa thread "
         "separada, para não bloquear o event loop com chamadas "
         "bloqueantes como psycopg2."),
        ("CORS", "mecanismo de segurança do navegador que controla quais "
         "origens podem consumir a API a partir do front-end; "
         "configurado aqui via CORSMiddleware."),
        ("422 vs 500", "convenção HTTP usada nesta API: 422 significa "
         "\"o parâmetro que você mandou é inválido\" (culpa de quem "
         "chamou), 500 significa \"algo falhou do nosso lado\" (banco, "
         "configuração) — distinção que ajuda quem consome a API a "
         "saber se deve corrigir a chamada ou tentar de novo mais tarde."),
    ]
    for termo, explicacao in conceitos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_page_break()


def build_etapa_4(doc: Document) -> None:
    doc.add_heading("Etapa 4 — Integração com API Externa", level=1)

    # -------------------------------------------------------------------
    doc.add_heading("1. Visão geral", level=2)
    doc.add_paragraph(
        "Até aqui, todo dado exibido vinha de dentro da própria "
        "operação (o banco de vendas). Esta etapa combina isso com uma "
        "fonte externa: a AwesomeAPI de câmbio (economia.awesomeapi.com."
        "br), pública e gratuita, sem necessidade de chave de API. A "
        "nova rota GET /vendas/resumo/cambio devolve o mesmo resumo de "
        "vendas de sempre, mas também convertido para a moeda pedida, "
        "usando a cotação obtida em tempo real."
    )
    doc.add_paragraph(
        "Diferente de uma consulta ao próprio banco, uma API externa "
        "introduz um tipo de risco novo: ela pode ficar lenta, fora do "
        "ar, ou mudar de resposta sem aviso. Boa parte do código desta "
        "etapa existe para lidar com isso, não para o \"caminho feliz\"."
    )

    # -------------------------------------------------------------------
    doc.add_heading("2. scripts/cambio.py: cache e fallback", level=2)
    add_code_block(
        doc,
        "def obter_cotacao(moeda: str) -> Cotacao:\n"
        "    moeda = normalizar_codigo_moeda(moeda)\n\n"
        "    with _cache_lock:\n"
        "        cotacao_cache = _cache.get(moeda)\n"
        "        if cotacao_cache and (time.time() - cotacao_cache.obtida_em) < TTL:\n"
        "            return replace(cotacao_cache, de_cache=True)\n\n"
        "    try:\n"
        "        cotacao = _buscar_na_api(moeda)\n"
        "    except (httpx.HTTPError, KeyError, ValueError) as exc:\n"
        "        with _cache_lock:\n"
        "            cotacao_cache = _cache.get(moeda)\n"
        "        if cotacao_cache is None:\n"
        "            raise CambioIndisponivel(...) from exc\n"
        "        return replace(cotacao_cache, de_cache=True, desatualizada=True)\n\n"
        "    with _cache_lock:\n"
        "        _cache[moeda] = cotacao\n"
        "    return cotacao",
    )
    doc.add_paragraph(
        "Três decisões concentradas nesta função, cada uma resolvendo um "
        "problema real de integração externa:"
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Cache com TTL: ")
    r.bold = True
    p.add_run(
        "CAMBIO_CACHE_TTL_SEGUNDOS (padrão 300s) evita bater na API "
        "externa a cada requisição -- mesmo sendo gratuita, ela tem "
        "limite de uso, e a cotação não muda segundo a segundo a ponto "
        "de justificar uma chamada de rede por request."
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Fallback com cache expirado: ")
    r.bold = True
    p.add_run(
        "se a API externa falhar (timeout, erro HTTP, resposta em "
        "formato inesperado) mas já existir uma cotação em cache — "
        "mesmo vencida —, ela é reaproveitada, marcada com "
        "desatualizada=True. Devolver um valor um pouco velho, sinalizado "
        "como tal, é melhor do que a rota inteira parar de funcionar só "
        "porque um serviço de terceiros teve um problema momentâneo. Só "
        "quando não há absolutamente nenhum valor em cache é que a "
        "função desiste, levantando CambioIndisponivel."
    )
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Lock de thread (_cache_lock): ")
    r.bold = True
    p.add_run(
        "como o FastAPI roda cada rota síncrona numa thread separada "
        "(decisão da Etapa 3), duas requisições podem chegar ao mesmo "
        "tempo e tentar ler/escrever o dicionário de cache "
        "simultaneamente. threading.Lock() garante que só uma thread "
        "mexe no cache por vez. replace() (do módulo dataclasses) cria "
        "uma cópia do objeto Cotacao em vez de alterar o que está no "
        "cache, para que threads concorrentes nunca enxerguem um objeto "
        "sendo modificado pela metade."
    )
    doc.add_paragraph(
        "Validação de formato (normalizar_codigo_moeda) acontece antes "
        "de qualquer acesso a cache ou rede -- um código como \"xx\" "
        "nunca chega a gerar uma chamada HTTP, é rejeitado na hora com "
        "ValueError."
    )

    # -------------------------------------------------------------------
    doc.add_heading("3. A rota (api/routers/vendas.py)", level=2)
    add_code_block(
        doc,
        "@router.get('/resumo/cambio', response_model=ResumoVendasCambio)\n"
        "def resumo_vendas_cambio(periodo: PeriodoDep, moeda: MoedaDep):\n"
        "    df = metricas.resumo_vendas_periodo(periodo.inicio, periodo.fim)\n"
        "    linha = df.iloc[0].to_dict()\n\n"
        "    cotacao = cambio.obter_cotacao(moeda)\n"
        "    fator = cotacao.valor or None\n\n"
        "    return {\n"
        "        ...\n"
        "        'faturamento_total_convertido': round(linha['faturamento_total'] / fator, 2),\n"
        "        'cotacao': {...},\n"
        "    }",
    )
    doc.add_paragraph(
        "moeda: MoedaDep segue exatamente o mesmo padrão de periodo: "
        "MoedaDep (dependencies.py) usa Depends(obter_moeda), que chama "
        "cambio.normalizar_codigo_moeda() e converte ValueError em "
        "HTTPException 422 -- a rota nunca vê um código de moeda "
        "inválido. cotacao.valor representa quantos reais valem 1 "
        "unidade da moeda estrangeira (ex.: 1 USD = 5,18 BRL), por isso "
        "a conversão de um valor em BRL para a moeda pedida é uma "
        "divisão (faturamento_brl / cotacao.valor), não uma "
        "multiplicação -- inverter essa conta é um erro fácil de "
        "cometer e passar despercebido, porque o número resultante ainda "
        "\"parece\" plausível à primeira vista."
    )

    # -------------------------------------------------------------------
    doc.add_heading("4. Erro da API externa: por que 503 e não 500", level=2)
    add_code_block(
        doc,
        "@app.exception_handler(CambioIndisponivel)\n"
        "def _cambio_indisponivel_handler(request, exc):\n"
        "    return JSONResponse(status_code=503, content={'detail': str(exc)})",
    )
    doc.add_paragraph(
        "Reaproveitando o padrão de exception_handler centralizado da "
        "Etapa 3, CambioIndisponivel ganhou seu próprio handler devolvendo "
        "503 (Service Unavailable), não 500. A distinção importa: 500 "
        "nesta API significa \"nosso código ou nosso banco tem um "
        "problema\"; 503 aqui significa \"um serviço de terceiros do qual "
        "dependemos está indisponível, o problema não é nosso, tente de "
        "novo mais tarde\". Quem consome a API pode reagir de forma "
        "diferente a cada um desses códigos."
    )

    # -------------------------------------------------------------------
    doc.add_heading("5. Script standalone (scripts/demo_cambio.py)", level=2)
    doc.add_paragraph(
        "Roda a mesma integração fora do contexto HTTP, para estudo e "
        "depuração isolada:"
    )
    add_code_block(doc, "python demo_cambio.py --periodo ultimos-30d --moeda USD")
    doc.add_paragraph(
        "Estruturalmente segue o mesmo padrão de gerar_relatorios.py "
        "(Etapa 2): argparse para os parâmetros, logging em vez de print "
        "para mensagens de status/erro, e os mesmos códigos de saída (0 "
        "sucesso, 1 erro de configuração/banco/câmbio, 2 argumento "
        "inválido). Importa metricas, cambio e periodo diretamente -- "
        "roda de dentro de scripts/, então não precisa da ponte de "
        "sys.path que a API usa (api/data_access.py)."
    )

    # -------------------------------------------------------------------
    doc.add_heading("6. O ciclo completo", level=2)
    doc.add_paragraph(
        "Com as quatro etapas prontas, uma chamada a "
        "GET /vendas/resumo/cambio?periodo=ultimos-30d&moeda=USD percorre "
        "toda a stack construída até aqui:"
    )
    ciclo = [
        "O cliente HTTP (navegador, curl, dashboard) chama a rota.",
        "FastAPI valida periodo e moeda via Depends() antes de executar "
        "qualquer lógica -- entrada inválida nunca chega ao passo seguinte.",
        "api/routers/vendas.py chama metricas.resumo_vendas_periodo(), que "
        "vive em scripts/ (Etapa 2).",
        "metricas.py chama database.executar_funcao(), que abre uma "
        "conexão com o Postgres (Neon/Supabase) e executa "
        "SELECT * FROM fn_resumo_vendas_periodo(...).",
        "A stored function (Etapa 1, PL/pgSQL) roda dentro do banco: "
        "valida o período, agrega vendas e itens_venda com JOIN, e devolve "
        "o resultado com RETURN QUERY.",
        "O resultado volta como linhas de cursor -> vira pandas.DataFrame "
        "em metricas.py -> vira dict em vendas.py.",
        "Em paralelo (mesma requisição), cambio.obter_cotacao() busca a "
        "cotação (cache ou API externa da AwesomeAPI, com fallback se ela "
        "falhar).",
        "A rota combina os dois resultados (faturamento em BRL + cotação) "
        "e devolve um dicionário Python.",
        "response_model=ResumoVendasCambio (Pydantic) valida esse "
        "dicionário e o serializa como JSON antes de enviar a resposta "
        "HTTP ao cliente.",
    ]
    for passo in ciclo:
        doc.add_paragraph(passo, style="List Number")
    doc.add_paragraph(
        "Cada camada depende só da camada logo abaixo dela, e cada uma "
        "tem uma responsabilidade que não se repete em nenhuma outra: o "
        "banco garante a regra de negócio (o que conta como faturamento), "
        "scripts/ sabe como conversar com o banco e com a API externa, e "
        "api/ só orquestra e expõe isso como contrato HTTP. É essa "
        "separação, mantida desde a Etapa 1, que permitiu adicionar a "
        "integração de câmbio nesta etapa tocando em poucos arquivos "
        "novos, sem reescrever nada das etapas anteriores."
    )

    # -------------------------------------------------------------------
    doc.add_heading("7. Conceitos-chave desta etapa", level=2)
    conceitos = [
        ("Timeout em chamada HTTP", "definir um tempo máximo de espera "
         "(CAMBIO_TIMEOUT_SEGUNDOS) para uma chamada de rede -- sem isso, "
         "uma API externa lenta prenderia a thread da requisição "
         "indefinidamente."),
        ("Cache com TTL (time to live)", "guardar um resultado por um "
         "tempo limitado para evitar chamadas repetidas desnecessárias, "
         "equilibrando \"dado atualizado\" com \"não sobrecarregar o "
         "serviço externo\"."),
        ("Fallback / degradação graciosa", "quando a fonte principal "
         "falha, usar a melhor alternativa disponível (cotação em cache, "
         "mesmo vencida) em vez de simplesmente quebrar -- sinalizando "
         "claramente que o dado está desatualizado."),
        ("threading.Lock", "mecanismo de sincronização que impede que "
         "duas threads modifiquem a mesma estrutura de dados (o dict de "
         "cache) ao mesmo tempo, evitando corrupção de estado."),
        ("dataclasses.replace", "cria uma cópia de um dataclass com "
         "alguns campos alterados, sem mutar o objeto original -- "
         "importante quando o objeto original é compartilhado entre "
         "threads (o cache)."),
        ("httpx.HTTPError", "classe-base de erro do httpx que cobre "
         "timeout, falha de conexão e respostas HTTP de erro "
         "(raise_for_status()) -- capturar essa classe cobre o conjunto "
         "inteiro de falhas de rede de uma vez."),
        ("Status code 503 vs 500", "503 comunica \"dependência externa "
         "indisponível\", diferente de 500 (\"erro interno\") -- "
         "distinção que ajuda quem consome a API a decidir se vale a "
         "pena tentar de novo."),
        ("Reuso entre CLI e API", "cambio.py e periodo.py são chamados "
         "tanto pelos scripts standalone quanto pelas rotas da API -- a "
         "mesma lógica, testada uma vez, é usada dos dois lugares."),
    ]
    for termo, explicacao in conceitos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_page_break()


def build_etapa_6(doc: Document) -> None:
    doc.add_heading("Etapa 6 — Ingestão e Conciliação de Dados", level=1)

    # -------------------------------------------------------------------
    doc.add_heading("1. O problema: chave surrogada não basta", level=2)
    doc.add_paragraph(
        "Até a Etapa 5, toda tabela usava só SERIAL (chave surrogada -- um "
        "número gerado pelo próprio banco, sem significado fora dele). Isso "
        "funciona bem para o modelo em si, mas quebra no momento em que "
        "chega uma carga de dados de fora: se um sistema externo reenviar "
        "informação sobre \"o mesmo cliente\" ou \"a mesma venda\", o banco "
        "não tem como saber que é o mesmo registro -- ele só enxerga IDs "
        "internos, que o sistema externo nunca viu. O resultado, sem uma "
        "chave natural, seria inserir tudo de novo a cada carga, duplicando "
        "clientes, vendas e itens."
    )
    doc.add_paragraph(
        "Chave natural é um identificador que já significa alguma coisa no "
        "mundo real (ou no sistema de origem): o SKU de um produto, o CPF "
        "de um cliente, o número de um pedido. Diferente da chave "
        "surrogada, ela existe independentemente do banco -- é o valor que "
        "permite reconhecer \"já vi isso antes\" ao receber uma carga nova. "
        "As duas convivem: a surrogada continua sendo a chave primária "
        "(mais estável e mais barata para JOIN), a natural vira uma "
        "constraint UNIQUE usada especificamente para reconciliação."
    )

    # -------------------------------------------------------------------
    doc.add_heading("2. Adicionando chaves a tabelas que já têm dados (sql/04)", level=2)
    add_code_block(
        doc,
        "ALTER TABLE produtos ADD COLUMN sku VARCHAR(30);\n\n"
        "UPDATE produtos\n"
        "SET sku = 'SKU-' || LPAD(id::text, 6, '0')\n"
        "WHERE sku IS NULL;\n\n"
        "ALTER TABLE produtos ALTER COLUMN sku SET NOT NULL;\n"
        "ALTER TABLE produtos ADD CONSTRAINT uq_produtos_sku UNIQUE (sku);",
    )
    doc.add_paragraph(
        "Não dá para simplesmente adicionar uma coluna NOT NULL numa tabela "
        "que já tem linhas -- o banco não saberia que valor colocar nas "
        "linhas existentes, e a instrução falharia. O padrão de migração "
        "usado aqui, e que vale para qualquer banco relacional em produção, "
        "tem 4 passos: (1) adiciona a coluna sem restrição nenhuma; "
        "(2) preenche (backfill) as linhas existentes com um valor válido; "
        "(3) só então marca a coluna como NOT NULL; (4) por fim, adiciona a "
        "constraint UNIQUE. Cada passo só é seguro porque o anterior já "
        "garantiu a condição que ele precisa."
    )
    doc.add_paragraph(
        "Para itens_venda.numero_linha (único por venda, não "
        "globalmente), o backfill é um pouco mais elaborado: usa "
        "ROW_NUMBER() OVER (PARTITION BY venda_id ORDER BY id) para "
        "numerar, dentro de cada venda, os itens existentes na ordem em "
        "que foram inseridos -- gerando um numero_linha (1, 2, 3...) "
        "válido para dados que nunca tiveram essa coluna."
    )

    # -------------------------------------------------------------------
    doc.add_heading("3. Staging / landing zone (sql/05)", level=2)
    doc.add_paragraph(
        "Staging (ou landing zone) é uma tabela que existe só para receber "
        "dados crus de uma fonte externa, exatamente como chegaram -- sem "
        "nenhuma regra do domínio aplicada ainda. stg_vendas_pedidos não "
        "tem FOREIGN KEY para clientes, produtos ou vendas de propósito: "
        "ela não valida se o cliente existe, se o produto existe, ou se o "
        "pedido é duplicado -- só recebe a linha. Essa é uma decisão de "
        "arquitetura deliberada, não um descuido: se a validação estivesse "
        "aqui, uma carga com um único registro problemático poderia travar "
        "a ingestão inteira. Separando staging de conciliação, o dado bruto "
        "sempre entra, e o processamento (que pode falhar por registro) "
        "acontece depois, sob controle."
    )
    doc.add_paragraph(
        "Reparo de design: nem toda coluna é NOT NULL em stg_vendas_"
        "pedidos. numero_pedido, numero_linha, cliente_cpf e produto_sku "
        "(as chaves usadas para buscar/reconciliar) são obrigatórias -- sem "
        "elas não há como processar a linha. Mas cliente_nome, "
        "cliente_email, produto_nome etc. são opcionais: um extrato real "
        "de um sistema externo pode vir com esses campos incompletos, e é "
        "exatamente esse cenário que a conciliação precisa saber tratar "
        "(seção 5)."
    )

    # -------------------------------------------------------------------
    doc.add_heading("4. A stored function de conciliação (sql/06)", level=2)
    doc.add_paragraph(
        "fn_conciliar_staging() percorre os pedidos pendentes em staging e, "
        "para cada um, tenta gravar cliente + venda + itens nas tabelas "
        "definitivas. Três mecanismos fazem o trabalho pesado:"
    )

    doc.add_heading("Lookup-ou-cria", level=3)
    add_code_block(
        doc,
        "SELECT id INTO v_cliente_id FROM clientes WHERE cpf = v_cliente_cpf;\n\n"
        "IF v_cliente_id IS NULL THEN\n"
        "    IF v_cliente_nome IS NULL OR trim(v_cliente_nome) = '' THEN\n"
        "        RAISE EXCEPTION 'cliente novo (cpf %) sem nome informado, requer revisao manual',\n"
        "            v_cliente_cpf;\n"
        "    END IF;\n"
        "    INSERT INTO clientes (...) VALUES (...) RETURNING id INTO v_cliente_id;\n"
        "END IF;",
    )
    doc.add_paragraph(
        "Busca pela chave natural (CPF); se não encontrar, cria um cliente "
        "novo -- mas só se houver dado suficiente (aqui, um nome). Sem "
        "isso, a função levanta uma exceção com uma mensagem explicando "
        "exatamente o que falta, em vez de criar um cadastro incompleto "
        "silenciosamente. A mesma lógica se repete para produto/SKU dentro "
        "do loop de itens."
    )

    doc.add_heading("Upsert idempotente via ON CONFLICT", level=3)
    add_code_block(
        doc,
        "INSERT INTO vendas (numero_pedido, cliente_id, data_venda, status)\n"
        "VALUES (v_pedido.numero_pedido, v_cliente_id, v_data_venda, v_status_venda)\n"
        "ON CONFLICT (numero_pedido) DO UPDATE\n"
        "    SET cliente_id = EXCLUDED.cliente_id,\n"
        "        data_venda = EXCLUDED.data_venda,\n"
        "        status     = EXCLUDED.status\n"
        "RETURNING id INTO v_venda_id;",
    )
    doc.add_paragraph(
        "ON CONFLICT (numero_pedido) DO UPDATE é o coração da idempotência: "
        "se já existe uma venda com esse numero_pedido (a constraint "
        "UNIQUE da Etapa 6), o Postgres atualiza a linha existente em vez "
        "de tentar inserir uma duplicata e falhar com erro de chave "
        "única. EXCLUDED é uma pseudo-tabela especial do PL/pgSQL/SQL que "
        "representa \"a linha que eu estava tentando inserir\" -- é como "
        "se referir aos valores do VALUES (...) de dentro do DO UPDATE. O "
        "mesmo padrão se repete no upsert de itens_venda, pela chave "
        "composta (venda_id, numero_linha)."
    )

    doc.add_heading("Isolamento de erro por pedido", level=3)
    add_code_block(
        doc,
        "FOR v_pedido IN SELECT DISTINCT numero_pedido FROM stg_vendas_pedidos WHERE ... LOOP\n"
        "    BEGIN\n"
        "        -- processa este pedido (cliente, venda, itens)\n"
        "        ...\n"
        "    EXCEPTION WHEN OTHERS THEN\n"
        "        -- desfaz so o que este pedido tentou gravar, e continua\n"
        "        UPDATE stg_vendas_pedidos SET status_processamento = 'erro', ...;\n"
        "        v_erros := v_erros + 1;\n"
        "    END;\n"
        "END LOOP;",
    )
    doc.add_paragraph(
        "Um bloco BEGIN...EXCEPTION...END aninhado dentro do loop cria, "
        "implicitamente, um savepoint por pedido. Se uma exceção acontece "
        "(seja o RAISE EXCEPTION manual, seja um erro de verdade do banco, "
        "como uma violação de constraint), tudo que aquele pedido tinha "
        "gravado até ali -- cliente novo, venda, itens -- é desfeito "
        "sozinho, sem afetar os pedidos já processados nem impedir que os "
        "próximos sejam tentados. Sem esse bloco, uma exceção não tratada "
        "abortaria a função inteira (e, com ela, tudo que já tinha sido "
        "processado na mesma chamada). Isso é o que permite uma carga com "
        "1.000 pedidos, dos quais 3 têm problema, terminar com 997 "
        "processados e 3 sinalizados para revisão -- em vez de nenhum."
    )

    doc.add_paragraph(
        "Nota de revisão: a primeira versão desta função contava "
        "\"ignorados\" (pedidos já processados numa chamada anterior) "
        "só olhando processado_em < inicio_da_chamada. O problema: quando "
        "a mesma carga é recebida de novo, o pedido É reprocessado nesta "
        "chamada (linhas novas 'pendente') E já tinha uma linha antiga "
        "'processado' -- as duas coisas seriam verdadeiras ao mesmo tempo, "
        "contando o mesmo pedido em \"atualizados\" e em \"ignorados\" "
        "simultaneamente. A correção foi guardar, num array "
        "(v_pedidos_tocados), todo numero_pedido que esta chamada de fato "
        "tentou processar, e excluir esses da contagem de ignorados no "
        "final. É um lembrete útil: mesmo lógica cuidadosamente escrita "
        "se beneficia de rastrear o cenário de ponta a ponta com números "
        "reais antes de confiar nela -- o que foi feito aqui comparando o "
        "resultado esperado de cada chamada de teste (seção 6) com uma "
        "simulação manual, statement a statement."
    )

    # -------------------------------------------------------------------
    doc.add_heading("5. Conceitos: idempotência e upsert", level=2)
    doc.add_paragraph(
        "Idempotência é a propriedade de uma operação que produz o mesmo "
        "resultado final não importa quantas vezes seja repetida com a "
        "mesma entrada. \"Inserir uma venda\" não é idempotente por si só "
        "(rodar duas vezes cria duas vendas); \"garantir que esta venda "
        "exista com estes dados\" é -- e é exatamente isso que o upsert "
        "(UPSERT = UPDATE + INSERT, via ON CONFLICT) implementa. Isso "
        "importa muito fora do ambiente controlado de um teste: arquivos "
        "de origem são reenviados por engano, jobs de carga falham na "
        "metade e são reexecutados do zero, uma fila de mensagens entrega "
        "o mesmo evento mais de uma vez -- um pipeline que não é "
        "idempotente duplica dado toda vez que um desses cenários (comuns) "
        "acontece."
    )

    # -------------------------------------------------------------------
    doc.add_heading("6. Carga incremental com watermark (CDC simplificado)", level=2)
    doc.add_paragraph(
        "O design atual (staging + conciliação) processa a carga inteira "
        "que está pendente a cada chamada -- não existe, hoje, um conceito "
        "de \"só buscar o que mudou desde a última vez\" na extração "
        "(fora do banco, antes de chegar em staging). Vale entender a "
        "técnica mesmo sem implementá-la aqui, porque é o próximo passo "
        "natural quando o volume de dados cresce."
    )
    doc.add_paragraph(
        "Watermark é um valor (tipicamente um timestamp, como "
        "updated_at) usado para marcar até onde a última carga já foi. "
        "Uma carga incremental funciona assim: guarda-se, em algum lugar "
        "de controle, o maior updated_at já processado com sucesso; a "
        "próxima extração busca só as linhas da origem com "
        "updated_at > watermark_salvo, processa essa fatia, e atualiza o "
        "watermark para o novo maior valor visto. Comparado a uma carga "
        "completa (reler a tabela de origem inteira toda vez), isso "
        "reduz drasticamente o volume transferido e processado a cada "
        "execução."
    )
    doc.add_paragraph(
        "Isso é uma forma simplificada de CDC (Change Data Capture) -- a "
        "categoria de técnicas para capturar apenas o que mudou numa "
        "fonte de dados. Um CDC \"de verdade\" (como o que ferramentas "
        "como Debezium implementam) lê o log de transações do banco de "
        "origem diretamente, capturando inclusive exclusões e sem "
        "depender de a aplicação de origem manter um updated_at correto. "
        "O watermark por coluna é a versão \"pobre\", mais simples de "
        "implementar sem infraestrutura extra, mas com uma limitação "
        "importante: se uma linha for fisicamente apagada na origem (não "
        "só atualizada), a carga baseada em updated_at nunca vai saber "
        "disso -- ela só enxerga o que ainda existe."
    )
    doc.add_paragraph(
        "Para aplicar isso neste projeto, a mudança seria: adicionar uma "
        "coluna updated_at (com DEFAULT NOW() e atualizada a cada UPDATE, "
        "via TRIGGER) nas tabelas de origem simuladas, guardar o "
        "watermark da última carga bem-sucedida numa tabela de controle "
        "(ex.: etl_controle(nome_carga, watermark)), e alterar a extração "
        "para filtrar por esse valor antes de gravar em staging."
    )

    # -------------------------------------------------------------------
    doc.add_heading("7. Como testar (Neon/Supabase)", level=2)
    passos = [
        "Rodar sql/04_chaves_naturais.sql (adiciona sku/cpf/numero_pedido/numero_linha).",
        "Rodar sql/05_staging.sql (cria stg_vendas_pedidos).",
        "Rodar sql/06_conciliacao.sql (cria fn_conciliar_staging).",
        "Rodar sql/07_staging_exemplo.sql (carga de exemplo: 3 pedidos).",
        "Chamar SELECT * FROM fn_conciliar_staging(); -> inseridos=2, atualizados=0, ignorados=0, erros=1.",
        "Recarregar sql/07_staging_exemplo.sql de novo (mesma carga) e chamar a função de novo "
        "-> inseridos=0, atualizados=2, ignorados=0, erros=1 -- e o total de linhas em vendas/"
        "itens_venda para esses pedidos NÃO cresce.",
        "Chamar a função uma terceira vez, sem recarregar staging "
        "-> inseridos=0, atualizados=0, ignorados=2, erros=0.",
    ]
    for passo in passos:
        doc.add_paragraph(passo, style="List Number")
    doc.add_paragraph(
        "Passo a passo com os comandos SQL exatos de cada checagem: "
        "README.md, seção \"Ingestão e conciliação de dados\"."
    )

    # -------------------------------------------------------------------
    doc.add_heading("8. Conceitos-chave desta etapa", level=2)
    conceitos = [
        ("Chave natural vs surrogada", "surrogada é um ID gerado pelo "
         "banco, sem significado fora dele; natural é um identificador que "
         "já existe no mundo real ou no sistema de origem (CPF, SKU, "
         "número de pedido) -- é ela que permite reconciliar cargas "
         "externas sem duplicar."),
        ("Migração com backfill", "padrão para adicionar uma coluna "
         "obrigatória a uma tabela que já tem dados: adicionar sem "
         "restrição, preencher os valores existentes, só então tornar "
         "obrigatória e única."),
        ("Staging / landing zone", "tabela que recebe dado bruto de uma "
         "fonte externa sem nenhuma validação de negócio, para que um "
         "registro problemático nunca impeça a ingestão dos demais."),
        ("Upsert (ON CONFLICT DO UPDATE)", "insere se a chave natural não "
         "existir, atualiza se existir -- numa única instrução atômica, "
         "sem race condition entre um SELECT de checagem e um INSERT "
         "separado."),
        ("EXCLUDED", "pseudo-tabela disponível dentro de ON CONFLICT DO "
         "UPDATE, representando a linha que estava sendo inserida quando "
         "o conflito ocorreu."),
        ("Idempotência", "propriedade de uma operação que, repetida com a "
         "mesma entrada, produz sempre o mesmo resultado final -- "
         "essencial em qualquer pipeline de ingestão, porque reenvio de "
         "dados e reprocessamento são a regra, não a exceção."),
        ("Savepoint implícito (BEGIN/EXCEPTION em PL/pgSQL)", "isola o "
         "efeito de um erro a um trecho específico do código, permitindo "
         "desfazer só aquela parte e continuar a execução, em vez de "
         "abortar a função inteira."),
        ("Watermark / carga incremental", "usar um valor de controle "
         "(tipicamente updated_at) para buscar, a cada execução, só o que "
         "mudou desde a última carga bem-sucedida, em vez de reler a "
         "origem inteira."),
        ("CDC (Change Data Capture)", "categoria de técnicas para "
         "capturar apenas as mudanças de uma fonte de dados; watermark por "
         "coluna é uma versão simplificada, que não captura exclusões "
         "físicas na origem."),
    ]
    for termo, explicacao in conceitos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_page_break()


def build_conclusao(doc: Document) -> None:
    doc.add_heading("Conclusão — Visão Consolidada do Projeto", level=1)

    # -------------------------------------------------------------------
    doc.add_heading("1. O projeto em uma tela", level=2)
    doc.add_paragraph(
        "Quatro etapas, quatro tecnologias, um único fluxo de dados. Este "
        "capítulo não introduz nada novo -- amarra o que já foi construído "
        "nas Etapas 1 a 4, para servir de revisão rápida antes de uma "
        "prova, entrevista técnica ou consulta futura."
    )
    add_code_block(
        doc,
        "PostgreSQL (Neon/Supabase)\n"
        "  |  tabelas: clientes, produtos, vendas, itens_venda\n"
        "  |  stored functions: fn_resumo_vendas_periodo, fn_top_produtos_faturamento,\n"
        "  |                    fn_faturamento_ticket_medio_cliente, fn_vendas_categoria_ranking\n"
        "  v\n"
        "scripts/database.py  --(psycopg2 + context manager)-->  scripts/metricas.py\n"
        "  |                                                        (SQL -> pandas.DataFrame)\n"
        "  v\n"
        "  +-- scripts/gerar_relatorios.py  --> CSV / JSON / resumo executivo (uso via cron)\n"
        "  |\n"
        "  +-- api/routers/*.py  --> api/schemas.py (Pydantic)  --> JSON (uso via HTTP)\n"
        "                |\n"
        "                +-- scripts/cambio.py --(httpx + cache/fallback)--> AwesomeAPI\n"
        "                        (só na rota /vendas/resumo/cambio e em demo_cambio.py)",
    )
    doc.add_paragraph(
        "Note que scripts/gerar_relatorios.py e api/ são dois \"consumidores\" "
        "diferentes da mesma camada de dados (scripts/metricas.py) -- nenhum "
        "dos dois duplica a lógica de conexão ou de query. Esse é o fio "
        "condutor de todo o projeto: cada regra existe em um único lugar."
    )

    # -------------------------------------------------------------------
    doc.add_heading("2. Checklist para rodar do zero", level=2)
    passos = [
        "Criar um projeto gratuito no Neon ou Supabase.",
        "No SQL Editor do serviço, rodar na ordem: 01_ddl.sql, 02_dml_seed.sql, 03_stored_procedures.sql.",
        "Copiar a connection string e colar em DATABASE_URL no .env (a partir de .env.example).",
        "Criar o ambiente virtual (python -m venv .venv) e ativar.",
        "Instalar as dependências (pip install -r requirements.txt).",
        "Testar a automação: python scripts/gerar_relatorios.py --periodo ultimos-30d.",
        "Testar a integração externa isolada: python scripts/demo_cambio.py --moeda USD.",
        "Subir a API a partir da raiz do projeto: uvicorn api.main:app --reload.",
        "Explorar as rotas pelo Swagger em http://127.0.0.1:8000/docs.",
    ]
    for passo in passos:
        doc.add_paragraph(passo, style="List Number")
    doc.add_paragraph(
        "Passo a passo completo, com os comandos exatos para Windows/"
        "PowerShell e exemplos de curl para cada rota: README.md na raiz do "
        "projeto."
    )

    # -------------------------------------------------------------------
    doc.add_heading("3. Glossário consolidado, por categoria", level=2)

    doc.add_heading("SQL e PL/pgSQL", level=3)
    for termo, explicacao in [
        ("SERIAL / FOREIGN KEY / CHECK", "mecanismos do próprio banco para "
         "garantir identidade, integridade referencial e validação de "
         "dados, independentemente de qual aplicação está escrevendo."),
        ("FUNCTION vs PROCEDURE (Postgres)", "FUNCTION devolve valor/tabela "
         "e é chamada com SELECT; PROCEDURE é chamada com CALL e não "
         "devolve result set do mesmo jeito -- por isso as 4 rotinas de "
         "métrica são FUNCTIONs com RETURN QUERY, não PROCEDUREs."),
        ("Window function (RANK)", "calcula um valor por linha (posição "
         "num ranking) considerando um conjunto relacionado de linhas, "
         "sem colapsar o resultado como GROUP BY faria."),
        ("generate_series / LATERAL", "recursos usados no seed para gerar "
         "milhares de linhas realistas a partir de SQL puro, sem precisar "
         "de um script externo."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_heading("Automação em Python", level=3)
    for termo, explicacao in [
        ("psycopg2 + context manager", "driver síncrono de Postgres, "
         "com a conexão gerenciada por um @contextmanager que garante "
         "fechamento mesmo em erro."),
        ("pandas.DataFrame", "estrutura usada para transformar o resultado "
         "de uma query em algo fácil de filtrar, agregar e exportar "
         "(to_csv, to_json)."),
        ("argparse + logging + exit codes", "base de todo script pensado "
         "para rodar sozinho via agendador: argumentos validados, "
         "mensagens com timestamp/nível, e um código de saída (0/1/2) que "
         "comunica sucesso ou falha a quem chamou o script."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_heading("FastAPI e Pydantic", level=3)
    for termo, explicacao in [
        ("response_model", "contrato de saída de uma rota: valida e "
         "serializa o retorno contra um schema Pydantic antes de mandar a "
         "resposta, e também gera a documentação automática."),
        ("Dependency injection (Depends)", "reaproveita lógica de "
         "validação (período, moeda) entre múltiplas rotas, escrita uma "
         "vez e injetada como parâmetro."),
        ("Rotas síncronas em threadpool", "rotas com def (não async def) "
         "rodam automaticamente numa thread separada -- necessário porque "
         "psycopg2 e httpx (uso síncrono) são bloqueantes."),
        ("exception_handler centralizado", "captura por tipo de exceção "
         "para toda a aplicação, evitando try/except repetido em cada "
         "rota, com status HTTP que comunica a natureza do erro (422 "
         "entrada inválida, 500 erro interno, 503 dependência externa "
         "fora do ar)."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_heading("Integração com API externa", level=3)
    for termo, explicacao in [
        ("Timeout", "tempo máximo de espera numa chamada de rede, para não "
         "prender a aplicação indefinidamente se o serviço externo travar."),
        ("Cache com TTL", "guarda um resultado por tempo limitado para "
         "reduzir chamadas repetidas a um serviço externo com limite de "
         "uso."),
        ("Fallback / degradação graciosa", "usar a última cotação "
         "conhecida (sinalizada como desatualizada) quando a fonte "
         "externa falha, em vez de quebrar a funcionalidade inteira."),
        ("httpx", "cliente HTTP usado tanto para consumir a API de câmbio "
         "quanto, internamente, pelo FastAPI TestClient nos testes."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{termo}: ")
        r.bold = True
        p.add_run(explicacao)

    # -------------------------------------------------------------------
    doc.add_heading("4. O que cada arquivo faz (referência rápida)", level=2)
    arquivos = [
        ("sql/01_ddl.sql", "cria as 4 tabelas e seus índices."),
        ("sql/02_dml_seed.sql", "popula clientes, produtos, vendas e itens_venda de exemplo."),
        ("sql/03_stored_procedures.sql", "cria as 4 stored functions de métricas."),
        ("sql/04_chaves_naturais.sql", "adiciona sku/cpf/numero_pedido/numero_linha, com backfill."),
        ("sql/05_staging.sql", "tabela de landing para cargas externas, sem regra de negócio."),
        ("sql/06_conciliacao.sql", "stored function que reconcilia staging com upsert idempotente."),
        ("sql/07_staging_exemplo.sql", "carga de exemplo para testar a conciliação."),
        ("scripts/database.py", "conexão com o Postgres e execução genérica de stored functions."),
        ("scripts/metricas.py", "uma função por stored function, devolvendo pandas.DataFrame."),
        ("scripts/periodo.py", "validação do parâmetro de período, usada pelo CLI e pela API."),
        ("scripts/cambio.py", "integração com a API de câmbio, com cache e fallback."),
        ("scripts/gerar_relatorios.py", "CLI que gera relatórios de vendas em CSV/JSON."),
        ("scripts/demo_cambio.py", "CLI que demonstra a integração de câmbio isolada."),
        ("api/main.py", "app FastAPI, CORS, handlers de erro, health check."),
        ("api/data_access.py", "ponte entre a API e os módulos de scripts/."),
        ("api/dependencies.py", "validação de periodo/moeda como dependency injection."),
        ("api/schemas.py", "modelos Pydantic das respostas da API."),
        ("api/routers/vendas.py", "rotas de resumo, top produtos, categoria e câmbio."),
        ("api/routers/clientes.py", "rota de ranking de clientes."),
    ]
    for nome, explicacao in arquivos:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{nome}: ")
        r.bold = True
        p.add_run(explicacao)

    doc.add_page_break()


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_titulo_capa(doc)
    build_etapa_1(doc)
    build_etapa_2(doc)
    build_etapa_3(doc)
    build_etapa_4(doc)
    build_etapa_6(doc)
    build_conclusao(doc)

    doc.save(OUTPUT_PATH)
    print(f"Documento gerado em: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
